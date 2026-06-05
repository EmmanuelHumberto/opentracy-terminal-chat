"""Conversores de documentos para Markdown (Fase 4).

Suporta:
  - .md / .txt  → copia direta
  - .pdf        → PyMuPDF + OCR (Tesseract) para paginas de imagem
  - .docx       → python-docx
  - .xlsx       → openpyxl
  - .jpg/.png/.bmp/.tiff → OCR com Tesseract
  - .zip        → extrai e converte arquivos suportados internamente

Todos os conversores retornam erros no schema padronizado (secao 12.3).
Casos de falha parcial retornam conversion_partial (recoverable: true).
"""

from __future__ import annotations

import os
import re
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Optional

from ligadoai_tools.safety import error_response


# ---------------------------------------------------------------------------
# Schema de erro padronizado
# ---------------------------------------------------------------------------

def _conversion_partial(detail: str = "") -> dict:
    return error_response(
        "conversion_partial",
        "Conversor extraiu parte do documento.",
        recoverable=True,
        detail=detail,
    )


def _conversion_failed(detail: str = "") -> dict:
    return error_response(
        "conversion_failed",
        "Conversor nao conseguiu processar o arquivo.",
        recoverable=False,
        detail=detail,
    )


# ---------------------------------------------------------------------------
# Variavel global para preservar estrutura de pastas
# ---------------------------------------------------------------------------

_current_source_dir: Optional[Path] = None


def _set_source_dir(source_dir: Optional[Path]) -> None:
    """Define o diretorio de origem atual para calcular caminhos relativos."""
    global _current_source_dir
    _current_source_dir = source_dir


def _output_path(source: Path, output_dir: Path, suffix: str = ".md") -> Path:
    """Calcula o caminho de saida preservando a estrutura relativa.

    Se source_dir estiver definido, mantem a estrutura de subpastas.
    Ex: source=knowledge/01-motores/x.md, source_dir=knowledge/
        -> output=knowledge_md/01-motores/x.md
    """
    global _current_source_dir
    if _current_source_dir and source.is_relative_to(_current_source_dir):
        rel = source.relative_to(_current_source_dir)
        if suffix:
            rel = rel.with_suffix(suffix)
        return output_dir / rel
    # Fallback: comportamento antigo (raiz)
    if suffix:
        return output_dir / f"{source.stem}{suffix}"
    return output_dir / source.name


# ---------------------------------------------------------------------------
# Utilitarios
# ---------------------------------------------------------------------------

def _safe_write(output_path: Path, content: str) -> None:
    """Escreve conteudo em arquivo, criando diretorios se necessario."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(content, encoding="utf-8")


def _markdown_header(source: str, title: str = "") -> str:
    """Retorna cabecalho padrao para arquivos convertidos."""
    lines = [
        f"<!-- Convertido de: {source} -->",
    ]
    if title:
        lines.append("")
        lines.append(f"# {title}")
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# .md / .txt (copia direta)
# ---------------------------------------------------------------------------

def convert_md(source: Path, output_dir: Path) -> dict[str, Any]:
    """Copia arquivo .md ou .txt para o diretorio de saida."""
    try:
        content = source.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return _conversion_failed(f"Arquivo nao e UTF-8 valido: {source.name}")
    except OSError as exc:
        return _conversion_failed(f"Erro de leitura: {exc}")

    output_path = _output_path(source, output_dir, suffix="")
    _safe_write(output_path, content)

    return {
        "success": True,
        "source": str(source),
        "output": str(output_path),
        "format": "md",
        "chars": len(content),
    }


# ---------------------------------------------------------------------------
# .jpg / .png / .bmp / .tiff (OCR com Tesseract)
# ---------------------------------------------------------------------------

def convert_image(source: Path, output_dir: Path) -> dict[str, Any]:
    """Converte imagem para Markdown usando OCR (Tesseract).

    Extrai texto de imagens como JPG, PNG, BMP, TIFF.
    Falha parcial: OCR nao conseguiu extrair texto (imagem sem texto legivel).
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return _conversion_failed(
            "pytesseract/Pillow nao instalado. Execute: pip install pytesseract Pillow"
        )

    try:
        img = Image.open(source)
    except Exception as exc:
        return _conversion_failed(f"Nao foi possivel abrir imagem: {exc}")

    try:
        ocr_text = pytesseract.image_to_string(img, lang="por+eng").strip()
    except Exception as exc:
        return _conversion_failed(f"Erro no OCR: {exc}")

    title = source.stem
    header = _markdown_header(str(source), title)

    if ocr_text:
        content = f"{header}\n\n## OCR da imagem\n\n{ocr_text}\n"
    else:
        content = f"{header}\n\n*[OCR nao extraiu texto desta imagem]*\n"

    output_path = _output_path(source, output_dir, suffix=".md")
    _safe_write(output_path, content)

    result: dict[str, Any] = {
        "success": True,
        "source": str(source),
        "output": str(output_path),
        "format": source.suffix.lower().lstrip("."),
        "chars": len(content),
        "ocr_text_chars": len(ocr_text),
    }

    if not ocr_text:
        result["warning"] = "OCR nao extraiu texto desta imagem."
        result["partial"] = _conversion_partial(
            "OCR nao conseguiu extrair texto da imagem. "
            "Pode ser uma imagem sem texto legivel ou com qualidade muito baixa."
        )

    return result


# ---------------------------------------------------------------------------
# .pdf (PyMuPDF + OCR)
# ---------------------------------------------------------------------------

def convert_pdf(source: Path, output_dir: Path) -> dict[str, Any]:
    """Converte PDF para Markdown usando PyMuPDF + OCR (Tesseract).

    Tenta extrair texto com PyMuPDF primeiro. Se a pagina estiver vazia
    (imagem escaneada), aplica OCR com Tesseract.
    """
    try:
        import fitz
    except ImportError:
        return _conversion_failed("pymupdf nao instalado. Execute: pip install pymupdf")

    _tesseract_available = False
    try:
        import pytesseract
        from PIL import Image
        import io
        _tesseract_available = True
    except ImportError:
        pass

    try:
        doc = fitz.open(source)
    except Exception as exc:
        return _conversion_failed(f"Nao foi possivel abrir PDF: {exc}")

    pages_content: list[str] = []
    empty_pages = 0
    ocr_pages = 0
    ocr_failed_pages = 0
    total_pages = len(doc)

    for page_num in range(total_pages):
        try:
            page = doc[page_num]
            text = page.get_text().strip()

            if text:
                pages_content.append(f"## Pagina {page_num + 1}\n\n{text}\n")
            elif _tesseract_available:
                ocr_pages += 1
                try:
                    pix = page.get_pixmap(dpi=300)
                    img_bytes = pix.tobytes("png")
                    img = Image.open(io.BytesIO(img_bytes))
                    ocr_text = pytesseract.image_to_string(img, lang="por+eng").strip()
                    if ocr_text:
                        pages_content.append(f"## Pagina {page_num + 1} (OCR)\n\n{ocr_text}\n")
                    else:
                        ocr_failed_pages += 1
                        empty_pages += 1
                        pages_content.append(f"*[pagina {page_num + 1} - OCR nao extraiu texto]*\n")
                except Exception as exc:
                    ocr_failed_pages += 1
                    empty_pages += 1
                    pages_content.append(f"*[pagina {page_num + 1} - erro OCR: {exc}]*\n")
            else:
                empty_pages += 1
                pages_content.append(f"*[pagina {page_num + 1} - sem texto extraivel]*\n")
        except Exception as exc:
            empty_pages += 1
            pages_content.append(f"*[pagina {page_num + 1} - erro: {exc}]*\n")

    doc.close()

    title = source.stem
    header = _markdown_header(str(source), title)
    content = header + "\n\n" + "\n\n".join(pages_content)

    output_path = _output_path(source, output_dir, suffix=".md")
    _safe_write(output_path, content)

    result: dict[str, Any] = {
        "success": True,
        "source": str(source),
        "output": str(output_path),
        "format": "pdf",
        "pages": total_pages,
        "chars": len(content),
    }

    if ocr_pages > 0:
        result["ocr_pages"] = ocr_pages
        result["ocr_failed"] = ocr_failed_pages

    if empty_pages > 0:
        result["warning"] = f"{empty_pages} de {total_pages} paginas sem texto"
        if ocr_pages > 0:
            result["warning"] += f" ({ocr_pages} OCR, {ocr_failed_pages} falhas)."
        result["partial"] = _conversion_partial(
            f"{empty_pages} pagina(s) sem texto. "
            f"{'OCR em ' + str(ocr_pages) + ' pag.' if ocr_pages > 0 else 'OCR indisponivel.'}"
        )

    return result


# ---------------------------------------------------------------------------
# .docx (python-docx)
# ---------------------------------------------------------------------------

def convert_docx(source: Path, output_dir: Path) -> dict[str, Any]:
    """Converte DOCX para Markdown."""
    try:
        from docx import Document
    except ImportError:
        return _conversion_failed("python-docx nao instalado.")

    try:
        doc = Document(str(source))
    except Exception as exc:
        return _conversion_failed(f"Erro ao abrir DOCX: {exc}")

    parts: list[str] = []
    nested_tables = 0

    def _process_paragraph(p: Any) -> str:
        text = p.text.strip()
        if not text:
            return ""
        style = p.style.name.lower() if p.style else ""
        if "heading 1" in style: return f"# {text}"
        elif "heading 2" in style: return f"## {text}"
        elif "heading 3" in style: return f"### {text}"
        elif "list" in style or "bullet" in style: return f"- {text}"
        return text

    def _process_table(table: Any, depth: int = 0) -> str:
        nonlocal nested_tables
        if depth > 0:
            nested_tables += 1
            return "*[tabela aninhada]*\n"
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        if not rows: return ""
        header_sep = "| " + " | ".join(["---"] * len(table.columns)) + " |"
        return "\n".join([rows[0], header_sep] + rows[1:]) + "\n"

    for para in doc.paragraphs:
        text = _process_paragraph(para)
        if text: parts.append(text)
    for table in doc.tables:
        table_md = _process_table(table)
        if table_md: parts.append(table_md)

    title = source.stem
    header = _markdown_header(str(source), title)
    content = header + "\n\n" + "\n\n".join(parts)

    output_path = _output_path(source, output_dir, suffix=".md")
    _safe_write(output_path, content)

    result = {"success": True, "source": str(source), "output": str(output_path), "format": "docx", "chars": len(content)}
    if nested_tables > 0:
        result["warning"] = f"{nested_tables} tabela(s) aninhada(s) achatadas."
        result["partial"] = _conversion_partial(f"{nested_tables} tabela(s) aninhada(s) achatada(s).")
    return result


# ---------------------------------------------------------------------------
# .xlsx (openpyxl)
# ---------------------------------------------------------------------------

def convert_xlsx(source: Path, output_dir: Path) -> dict[str, Any]:
    """Converte XLSX para Markdown."""
    try:
        import openpyxl
    except ImportError:
        return _conversion_failed("openpyxl nao instalado.")

    try:
        wb = openpyxl.load_workbook(source, read_only=True, data_only=True)
    except Exception as exc:
        return _conversion_failed(f"Erro ao abrir XLSX: {exc}")

    merged_cells_count = 0
    parts: list[str] = []
    title = source.stem
    header = _markdown_header(str(source), title)
    parts.append(header)
    parts.append("")

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"## Planilha: {sheet_name}\n")
        merged = set()
        for merge_range in ws.merged_cells.ranges:
            merged_cells_count += 1
            for row in range(merge_range.min_row, merge_range.max_row + 1):
                for col in range(merge_range.min_col, merge_range.max_col + 1):
                    merged.add((row, col))
        rows_data = []
        for row in ws.iter_rows(min_row=1, values_only=False):
            row_values = []
            for cell in row:
                val = "" if (cell.row, cell.column) in merged else cell.value
                if val is None: val = ""
                else: val = str(val).replace("\n", " ").strip()
                row_values.append(val)
            if any(v for v in row_values): rows_data.append(row_values)
        if rows_data:
            headers = rows_data[0]
            table_lines = []
            table_lines.append("| " + " | ".join(headers) + " |")
            table_lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in rows_data[1:]:
                while len(row) < len(headers): row.append("")
                table_lines.append("| " + " | ".join(row) + " |")
            parts.append("\n".join(table_lines))
        else:
            parts.append("*[planilha vazia]*")
        parts.append("")

    wb.close()
    content = "\n".join(parts)

    output_path = _output_path(source, output_dir, suffix=".md")
    _safe_write(output_path, content)

    result = {"success": True, "source": str(source), "output": str(output_path), "format": "xlsx", "chars": len(content)}
    if merged_cells_count > 0:
        result["warning"] = f"{merged_cells_count} celula(s) mesclada(s)."
        result["partial"] = _conversion_partial(f"{merged_cells_count} celula(s) mesclada(s).")
    return result


# ---------------------------------------------------------------------------
# .zip (extrai e converte arquivos internos)
# ---------------------------------------------------------------------------

def convert_zip(source: Path, output_dir: Path) -> dict[str, Any]:
    """Extrai um arquivo ZIP e converte todos os arquivos suportados.

    Para cada arquivo dentro do ZIP com extensao suportada, aplica o
    conversor correspondente e salva o Markdown no diretorio de saida.

    Retorna:
      - success: True se pelo menos um arquivo foi convertido
      - converted_files: lista de resultados individuais
      - errors_files: lista de erros internos
      - unsupported: lista de arquivos ignorados (extensao nao suportada)
    """
    if not source.is_file():
        return _conversion_failed(f"Arquivo ZIP nao encontrado: {source}")

    try:
        zf = zipfile.ZipFile(source, 'r')
    except zipfile.BadZipFile:
        return _conversion_failed(f"Arquivo ZIP invalido ou corrompido: {source.name}")
    except Exception as exc:
        return _conversion_failed(f"Erro ao abrir ZIP: {exc}")

    zip_name = source.stem
    converted_files: list[dict[str, Any]] = []
    errors_files: list[dict[str, Any]] = []
    unsupported: list[str] = []
    total_chars = 0

    with tempfile.TemporaryDirectory(prefix=f"zip_{zip_name}_") as tmp_dir:
        tmp_path = Path(tmp_dir)

        for info in zf.infolist():
            if info.is_dir():
                continue

            filename = info.filename
            ext = Path(filename).suffix.lower()

            if "/__MACOSX" in filename or filename.startswith("__MACOSX"):
                continue
            if "/." in filename or filename.startswith("."):
                continue

            if ext not in SUPPORTED_EXTENSIONS:
                unsupported.append(filename)
                continue

            try:
                zf.extract(info, tmp_path)
            except Exception as exc:
                errors_files.append({
                    "file": filename,
                    "error": str(exc),
                })
                continue

            extracted = tmp_path / filename

            if not extracted.is_file():
                errors_files.append({
                    "file": filename,
                    "error": "Arquivo extraido nao encontrado",
                })
                continue

            try:
                result = SUPPORTED_EXTENSIONS[ext](extracted, output_dir)
                if result.get("success"):
                    converted_files.append(result)
                    total_chars += result.get("chars", 0)
                else:
                    errors_files.append({
                        "file": filename,
                        "error": result.get("error", {}).get("message", "Falha na conversao"),
                    })
            except Exception as exc:
                errors_files.append({
                    "file": filename,
                    "error": str(exc),
                })

    zf.close()

    result: dict[str, Any] = {
        "success": len(converted_files) > 0,
        "source": str(source),
        "output": str(output_dir),
        "format": "zip",
        "converted_files": converted_files,
        "converted_count": len(converted_files),
        "errors_count": len(errors_files),
        "unsupported_count": len(unsupported),
        "total_chars": total_chars,
    }

    if errors_files:
        result["errors_files"] = errors_files
    if unsupported:
        result["unsupported_files"] = unsupported
    if len(converted_files) == 0 and len(errors_files) > 0:
        result["partial"] = _conversion_partial(
            f"Nenhum arquivo convertido. {len(errors_files)} erro(s), "
            f"{len(unsupported)} arquivo(s) ignorados."
        )

    return result


# ---------------------------------------------------------------------------
# Registro de extensoes e funcoes publicas
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: dict[str, Any] = {
    ".md": convert_md,
    ".txt": convert_md,
    ".pdf": convert_pdf,
    ".docx": convert_docx,
    ".xlsx": convert_xlsx,
    ".jpg": convert_image,
    ".jpeg": convert_image,
    ".png": convert_image,
    ".bmp": convert_image,
    ".tiff": convert_image,
    ".tif": convert_image,
    ".zip": convert_zip,
}


def convert_file(source: Path, output_dir: Path) -> dict[str, Any]:
    """Converte um arquivo para Markdown."""
    ext = source.suffix.lower()
    converter = SUPPORTED_EXTENSIONS.get(ext)
    if not converter:
        return _conversion_failed(f"Formato nao suportado: {ext}")
    return converter(source, output_dir)


def convert_directory(
    source_dir: Path, output_dir: Path, *, recursive: bool = True
) -> dict[str, Any]:
    """Converte todos os arquivos suportados em um diretorio.

    Preserva a estrutura de subpastas: knowledge/01-motores/x.md
    → knowledge_md/01-motores/x.md
    """
    if not source_dir.is_dir():
        return _conversion_failed(f"Diretorio nao encontrado: {source_dir}")

    _set_source_dir(source_dir)

    results: list[dict[str, Any]] = []
    errors: list[dict[str, Any]] = []
    partials: list[dict[str, Any]] = []
    total_chars = 0

    pattern = "**/*" if recursive else "*"
    for entry in sorted(source_dir.glob(pattern)):
        if not entry.is_file():
            continue
        ext = entry.suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            continue
        result = convert_file(entry, output_dir)
        if result.get("success"):
            results.append(result)
            total_chars += result.get("chars", 0)
        elif "error" in result:
            errors.append(result)
        if result.get("partial"):
            partials.append(result)

    _set_source_dir(None)

    return {
        "success": len(errors) == 0,
        "total_files": len(results) + len(errors),
        "converted": len(results),
        "errors": len(errors),
        "partials": len(partials),
        "total_chars": total_chars,
        "source_dir": str(source_dir),
        "output_dir": str(output_dir),
        "results": results,
        "error_details": errors,
        "partial_details": partials,
    }
